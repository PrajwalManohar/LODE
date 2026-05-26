"""Word-document renderer for LODE custom SOPs.

Content (sample notes, parameters, procedure, etc.) is built by the
`vein.services.sop_builder` module using the LLM + RAG. This file is the
*pure renderer* — it lays out the structured content into a styled .docx
that matches the reference SOP screenshot:

    Header card
    Customized sample-specific notes (callout)
    Recommended operating parameters (2-col table)
    Pre-session checklist
    Step-by-step procedure (numbered)
    Post-run checklist
    Safety notes
    RAG references
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from vein.config import OUTPUT_DIR, ensure_dirs, local_tz_label
from vein.models.experiment import BookingOption, ExperimentContext, InstrumentFit
from vein.services.sop_builder import build_sop_content

NAVY = RGBColor(0x21, 0x31, 0x4F)
GOLD = RGBColor(0xC7, 0x9A, 0x3E)
INK = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
CITE = RGBColor(0x3B, 0x6F, 0xB5)
LIGHT_BG = "F4F5F7"
HEADER_BG = "21314F"
CALLOUT_BG = "FFF7E6"
TABLE_HEADER_BG = "E5E7EB"


def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "D1D5DB") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def _add_run(paragraph, text: str, *, bold=False, color=None, size=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def _add_citation(paragraph, citation: str) -> None:
    if not citation:
        return
    paragraph.add_run("  ")
    _add_run(paragraph, f"[{citation}]", italic=True, color=CITE, size=9)


def _header_card(doc, *, booking_code: str, instrument: str, ctx: ExperimentContext,
                 booking: BookingOption, fit: InstrumentFit) -> None:
    """Two-row card: dark band with title; light band with key facts."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.autofit = True
    tbl.allow_autofit = True

    # Top — dark band
    top = tbl.rows[0].cells[0]
    _shade(top, HEADER_BG)
    p = top.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, "LODE  ", bold=True, color=GOLD, size=16)
    _add_run(p, "· Custom Session SOP", color=RGBColor(0xFF, 0xFF, 0xFF), size=12)
    p2 = top.add_paragraph()
    _add_run(p2, f"{instrument} — {ctx.analysis_goal or 'session'}",
             bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=14)
    p3 = top.add_paragraph()
    _add_run(p3, f"Booking {booking_code}   ·   Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}",
             color=RGBColor(0xCB, 0xD5, 0xE1), size=9)

    # Bottom — facts grid
    bot = tbl.rows[1].cells[0]
    _shade(bot, LIGHT_BG)
    bot.paragraphs[0].text = ""

    inner = bot.add_table(rows=2, cols=4)
    inner.autofit = True
    items = [
        ("INSTRUMENT", instrument),
        ("RESEARCHER", ctx.researcher_name or "—"),
        ("RESEARCH GROUP", ctx.research_group or "—"),
        ("MATERIAL", ctx.material_type or "—"),
        ("SAMPLE", ctx.sample_dimensions or "—"),
        ("SESSION", f"{booking.start_time:%a, %b %d · %H:%M}–{booking.end_time:%H:%M} {local_tz_label()}"),
        ("PREP STARTS", booking.prep_start.strftime("%a, %b %d · %H:%M")),
        ("FIT SCORE", f"{fit.fit_score}/100 ({fit.grade})"),
    ]
    for i, (k, v) in enumerate(items):
        cell = inner.cell(i // 4, i % 4)
        cell.text = ""
        p = cell.paragraphs[0]
        _add_run(p, k, bold=True, color=GRAY, size=8)
        p2 = cell.add_paragraph()
        _add_run(p2, v, color=INK, size=10)


def _section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, text.upper(), bold=True, color=NAVY, size=11)


def _callout(doc, text: str, *, label: str = "Customized sample-specific notes") -> None:
    if not text:
        return
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _shade(cell, CALLOUT_BG)
    p = cell.paragraphs[0]
    _add_run(p, label.upper(), bold=True, color=GOLD, size=8)
    p2 = cell.add_paragraph()
    _add_run(p2, text, color=INK, size=10)


def _parameter_table(doc, params: list[dict]) -> None:
    if not params:
        return
    tbl = doc.add_table(rows=1 + len(params), cols=2)
    tbl.style = "Light List"

    # header
    hdr_p = tbl.rows[0].cells[0].paragraphs[0]
    _add_run(hdr_p, "Parameter", bold=True, color=INK, size=10)
    _shade(tbl.rows[0].cells[0], TABLE_HEADER_BG)
    hdr_v = tbl.rows[0].cells[1].paragraphs[0]
    _add_run(hdr_v, "Recommended value (with citation)", bold=True, color=INK, size=10)
    _shade(tbl.rows[0].cells[1], TABLE_HEADER_BG)

    for i, p in enumerate(params, start=1):
        label_cell = tbl.rows[i].cells[0]
        value_cell = tbl.rows[i].cells[1]
        _set_cell_borders(label_cell); _set_cell_borders(value_cell)
        _add_run(label_cell.paragraphs[0], p.get("label", ""), bold=True, color=INK, size=10)
        vp = value_cell.paragraphs[0]
        _add_run(vp, p.get("value", ""), color=INK, size=10)
        _add_citation(vp, p.get("citation", ""))


def _bullet(doc, item: dict, *, bullet_char: str = "☐") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    _add_run(p, f"{bullet_char}  ", color=NAVY, size=11)
    _add_run(p, item.get("text", ""), color=INK, size=10)
    _add_citation(p, item.get("citation", ""))


def _numbered_step(doc, n: int, step: dict) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.05)
    _add_run(p, f"{n}. ", bold=True, color=NAVY, size=11)
    _add_run(p, step.get("title", "").strip(), bold=True, color=INK, size=10)
    if step.get("detail"):
        _add_run(p, " — ", color=GRAY, size=10)
        _add_run(p, step["detail"].strip(), color=INK, size=10)
    _add_citation(p, step.get("citation", ""))


def _references_section(doc, refs: list[dict]) -> None:
    if not refs:
        return
    _section_heading(doc, "RAG References")
    intro = doc.add_paragraph()
    _add_run(intro,
             "This SOP was generated by LODE's authoring agent grounded in the following "
             "lab knowledge sources (manuals, lab SOPs, maintenance logs).",
             color=GRAY, size=9, italic=True)
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        _add_run(p, f"[{i}] ", bold=True, color=NAVY, size=10)
        _add_run(p, r.get("label", ""), color=INK, size=10)


# ---------------------------------------------------------------------------
# Public entry point — used by vein.agents.pipeline.confirm_booking
# ---------------------------------------------------------------------------
def generate_sop_document(
    ctx: ExperimentContext,
    fit: InstrumentFit,
    booking: BookingOption,
    *,
    booking_code: str = "",
) -> Path:
    ensure_dirs()
    content = build_sop_content(ctx, fit, booking)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Tighten margins so the parameter table doesn't wrap awkwardly
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.7)

    _header_card(
        doc,
        booking_code=booking_code or "—",
        instrument=fit.instrument_name,
        ctx=ctx,
        booking=booking,
        fit=fit,
    )

    if content.get("sample_notes"):
        _callout(doc, content["sample_notes"])

    _section_heading(doc, "Recommended Operating Parameters")
    _parameter_table(doc, content.get("parameters", []))

    _section_heading(doc, "Pre-Session Checklist  (customized to sample)")
    for item in content.get("pre_checklist", []):
        _bullet(doc, item, bullet_char="☐")

    _section_heading(doc, "Step-by-Step Procedure")
    for i, step in enumerate(content.get("procedure", []), 1):
        _numbered_step(doc, i, step)

    _section_heading(doc, "Post-Run Checklist")
    for item in content.get("post_checklist", []):
        _bullet(doc, item, bullet_char="☐")

    if content.get("safety_notes"):
        _section_heading(doc, "Safety Notes")
        for item in content["safety_notes"]:
            _bullet(doc, item, bullet_char="•")

    _references_section(doc, content.get("references", []))

    # Fit rationale footnote
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    _add_run(p, "Fit rationale: ", bold=True, color=GRAY, size=9)
    _add_run(p, fit.rationale, italic=True, color=GRAY, size=9)

    fname = f"sop_{fit.instrument_id}_{booking.start_time.strftime('%Y%m%d_%H%M')}.docx"
    path = OUTPUT_DIR / fname
    doc.save(path)
    return path
