"""Fire any of the four LODE automation emails with realistic sample data.

Use this to demo the automation flow without having to walk through the full UI
booking pipeline every time. Each scenario calls the real sender, so it goes
through the configured transport (Resend → SendGrid → SMTP → local outbox) and
records an ``automation_event`` row exactly like a real booking would.

Usage
-----
    python scripts/demo_emails.py <scenario> [--to <email>]

Scenarios
---------
    booking          Email 1 · Booking confirmation + SOP attachment (navy)
    hitl-training    Email 2 · HITL — researcher lacks SEM-Operator training (brown)
    hitl-hazmat      Email 2 · HITL — hydrofluoric acid detected (brown)
    hitl-confidence  Email 2 · HITL — fit confidence below 80%        (brown)
    workorder-preemptive  Email 3 · Calibration overdue at booking    (purple)
    workorder-postrun     Email 3 · Detector saturation in post-run   (purple)
    monthly          Email 4 · Monthly utilization report             (green)
    all              Fire all of the above, in order

Flags
-----
    --to <email>     Override every recipient with this address (handy for
                     Resend test mode — see EMAIL_OVERRIDE in .env).
    --list           Print the scenario list and exit.
"""

from __future__ import annotations

import argparse
import os
import sys
from datetime import datetime, timedelta
from pathlib import Path

# Allow running as `python scripts/demo_emails.py ...` from the project root.
ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from vein.config import LOCAL_TZ_NAME, local_now, local_tz_label, settings  # noqa: E402
from vein.models.experiment import ExperimentContext  # noqa: E402
from vein.services import email as email_service  # noqa: E402


# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def _hr(title: str) -> None:
    bar = "─" * max(8, 78 - len(title))
    print(f"\n  ▌ {title} {bar}\n")


def _summarize(result: dict, label: str) -> None:
    print(f"    ✓ {label}")
    print(f"      transport : {result.get('transport')}")
    print(f"      delivered : {result.get('sent')}")
    print(f"      to        : {', '.join(result.get('to') or [])}")
    print(f"      record_id : {result.get('id')}")


def _ensure_sop(researcher: str, booking_code: str) -> Path:
    """A real SOP exists in data/output/sops after a real booking. For the
    standalone demo we synthesize a tiny docx so the attachment is present."""
    out_dir = ROOT / "data" / "output" / "sops"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"demo_sop_{booking_code}.docx"
    if not path.exists():
        try:
            from docx import Document  # python-docx

            doc = Document()
            doc.add_heading(f"LODE — Demo SOP ({booking_code})", level=1)
            doc.add_paragraph(f"Researcher: {researcher}")
            doc.add_paragraph(
                "This is a synthetic SOP attached by scripts/demo_emails.py to "
                "demonstrate the email automation. A real booking generates a "
                "fully cited multi-section document via vein.services.sop_docx."
            )
            doc.save(path)
        except Exception:
            # Last-resort plaintext stub if python-docx isn't importable
            path.write_bytes(b"LODE demo SOP placeholder.")
    return path


# ----------------------------------------------------------------------------
# Scenario 1 — Email 1 · Booking confirmation + SOP
# ----------------------------------------------------------------------------
def scenario_booking(to_override: str | None) -> None:
    _hr("Email 1 — Booking confirmation + SOP  (navy)")
    booking_code = "VEIN-9001"
    start = (local_now() + timedelta(days=2)).replace(hour=9, minute=0, second=0, microsecond=0)
    end = start + timedelta(hours=3)
    ctx = ExperimentContext(
        material_type="martensitic steel",
        analysis_goal="fracture surface morphology",
        sample_dimensions="5mm × 5mm",
        surface_condition="uncoated fracture surface",
        coating_status="coating scheduled",
        urgency="high",
        deadline="Thursday",
        researcher_name="Dr. Sarah Chen",
        researcher_email=to_override or "chen@mines.edu",
        research_group="MetEng-Lab",
        trained_instruments=["SEM-Operator", "XRD-Safety-101"],
        notes="Demo run from scripts/demo_emails.py",
        is_complete=True,
    )
    sop_path = _ensure_sop(ctx.researcher_name, booking_code)
    tz_label = local_tz_label()
    when_str = f"{start:%A, %b %d, %Y · %I:%M %p}–{end:%I:%M %p} {tz_label}"
    checklist = [
        {
            "text": "Carbon coat sample at Station 2B (90 min prep).",
            "citation": "Mines SEM Lab SOP §2.1, p.8",
        },
        {
            "text": "Warm up JEOL column — minimum 25 min idle vacuum.",
            "citation": "JEOL JSM-IT800 Manual §3.2, p.42",
        },
        {
            "text": "Load specimen at 10 mm WD, 15 kV, spot size 50 — EDS preset.",
            "citation": "JEOL JSM-IT800 Manual §4.4, p.61",
        },
    ]
    ics_text = (
        "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//LODE//Lab//EN\r\nBEGIN:VEVENT\r\n"
        f"UID:{booking_code}@lode\r\nDTSTART:{start.strftime('%Y%m%dT%H%M%S')}\r\n"
        f"DTEND:{end.strftime('%Y%m%dT%H%M%S')}\r\nSUMMARY:LODE session — JEOL SEM-EDS\r\n"
        "LOCATION:Brown Hall B12\r\nDESCRIPTION:Demo session via demo_emails.py\r\n"
        "END:VEVENT\r\nEND:VCALENDAR"
    )
    result = email_service.send_sop_email(
        ctx,
        str(sop_path),
        booking_code=booking_code,
        instrument="JEOL JSM-IT800 SEM-EDS",
        location="Brown Hall B12",
        when=when_str,
        fit_score=92,
        grade="A",
        approved_by="Auto-approved · LODE safety gate",
        checklist=checklist,
        ics_text=ics_text,
    )
    _summarize(result, "Email 1 dispatched")


# ----------------------------------------------------------------------------
# Scenario 2 — Email 2 · HITL approval requests (three variants)
# ----------------------------------------------------------------------------
def _send_hitl(
    *,
    booking_code: str,
    researcher: str,
    alert_title: str,
    alert_text: str,
    reasoning: list[str],
    confidence: int,
    training_status: str,
    to_override: str | None,
) -> None:
    if to_override:
        # The HITL sender writes to settings.lab_email_tech; honor an override
        # by mutating it transiently so the dispatch routes correctly.
        original = settings.lab_email_tech
        settings.lab_email_tech = to_override
    try:
        start = (local_now() + timedelta(days=3)).replace(hour=14, minute=0, second=0, microsecond=0)
        end = start + timedelta(hours=3)
        when_str = f"{start:%A, %b %d, %Y · %I:%M %p}–{end:%I:%M %p} {local_tz_label()}"
        result = email_service.send_hitl_email(
            booking_code=booking_code,
            researcher=researcher,
            instrument="JEOL JSM-IT800 SEM-EDS",
            location="Brown Hall B12",
            when=when_str,
            experiment="Hydrogen permeation — fracture surface morphology",
            fit_score=92,
            grade="A",
            confidence=confidence,
            training_status=training_status,
            alert_title=alert_title,
            alert_text=alert_text,
            reasoning=reasoning,
        )
        _summarize(result, f"Email 2 dispatched ({alert_title})")
    finally:
        if to_override:
            settings.lab_email_tech = original


def scenario_hitl_training(to_override: str | None) -> None:
    _hr("Email 2 — HITL: missing training  (brown)")
    _send_hitl(
        booking_code="VEIN-9002",
        researcher="Dr. Sarah Chen (chen@mines.edu)",
        alert_title="Training certification missing",
        alert_text=(
            "Researcher lacks documented SEM-Operator training, "
            "required for JEOL JSM-IT800 SEM-EDS per safety SOP §1.4."
        ),
        reasoning=[
            "Agent 1 parsed material = martensitic steel, goal = morphology characterization.",
            "Agent 2 ranked JEOL SEM-EDS at fit 92 (Grade A); confidence 88%.",
            "Safety gate matched required_training=['SEM-Operator'] against researcher profile — no record found.",
            "Calibration hours: 187h since last cal (within 400h limit — OK).",
            "Refusal rule #1 (missing training) triggered → routed to HITL queue.",
        ],
        confidence=88,
        training_status="SEM-Operator: NOT ON FILE",
        to_override=to_override,
    )


def scenario_hitl_hazmat(to_override: str | None) -> None:
    _hr("Email 2 — HITL: hazardous materials  (brown)")
    _send_hitl(
        booking_code="VEIN-9003",
        researcher="Dr. Aiyana Bennett (bennett@mines.edu)",
        alert_title="Hazardous materials detected",
        alert_text=(
            "Experiment description references hydrofluoric acid — EH&S review "
            "required before booking can proceed."
        ),
        reasoning=[
            "Agent 1 parsed material = mine drainage water, goal = trace metal quantification.",
            "Safety gate matched HAZMAT_KEYWORDS: 'hydrofluoric' (HF acid).",
            "Refusal rule #3 (hazmat) triggered — manual EH&S sign-off required.",
            "Suggested next step: confirm sample digestion uses HF only in EH&S-approved hood.",
        ],
        confidence=91,
        training_status="ICP-MS-Cert: ON FILE (since 2025-09-12)",
        to_override=to_override,
    )


def scenario_hitl_confidence(to_override: str | None) -> None:
    _hr("Email 2 — HITL: fit confidence below 80%  (brown)")
    _send_hitl(
        booking_code="VEIN-9004",
        researcher="Marcus Rivera (marcus@mines.edu)",
        alert_title="Fit Scorer confidence below 80%",
        alert_text=(
            "The Fit Scorer recommended an instrument with confidence 64% — "
            "below the 80% architectural floor. A human reviewer should "
            "confirm whether this instrument is appropriate before scheduling."
        ),
        reasoning=[
            "Agent 1 parsed material = unknown amorphous powder.",
            "Agent 2 best match: XRD at fit 62 (Grade C); confidence 64%.",
            "Other instruments scored ≤45 — no high-confidence recommendation.",
            "Refusal rule #4 (confidence floor) triggered → routed to HITL queue.",
        ],
        confidence=64,
        training_status="XRD-Safety-101: ON FILE",
        to_override=to_override,
    )


# ----------------------------------------------------------------------------
# Scenario 3 — Email 3 · Maintenance work order (two variants)
# ----------------------------------------------------------------------------
def _send_workorder(
    *,
    code: str,
    instrument: str,
    location: str,
    usage_hours: float,
    interval_hours: float,
    last_calibrated: str,
    issue_type: str,
    severity: str,
    triggered_by: str,
    anomaly: str,
    actions: list[str],
    to_override: str | None,
) -> None:
    if to_override:
        orig_tech = settings.lab_email_tech
        orig_fac = settings.lab_email_facilities
        settings.lab_email_tech = to_override
        settings.lab_email_facilities = to_override
    try:
        result = email_service.send_work_order_email(
            work_order_code=code,
            instrument=instrument,
            location=location,
            usage_hours=usage_hours,
            interval_hours=interval_hours,
            last_calibrated=last_calibrated,
            issue_type=issue_type,
            severity=severity,
            triggered_by=triggered_by,
            anomaly=anomaly,
            actions=actions,
        )
        _summarize(result, f"Email 3 dispatched ({code} · {severity})")
    finally:
        if to_override:
            settings.lab_email_tech = orig_tech
            settings.lab_email_facilities = orig_fac


def scenario_workorder_preemptive(to_override: str | None) -> None:
    _hr("Email 3 — Work order: calibration overdue at booking  (purple)")
    _send_workorder(
        code="WO-001",
        instrument="JEOL JSM-IT800 SEM-EDS",
        location="Brown Hall B12",
        usage_hours=315.0,
        interval_hours=250.0,
        last_calibrated="2026-02-14",
        issue_type="Calibration overdue",
        severity="Warning",
        triggered_by="Safety gate · pre-booking check (BKG-2026-0142)",
        anomaly="Usage 315h logged vs 250h calibration interval (126%).",
        actions=[
            "Run NIST 612 standard for energy calibration before next session "
            "[Source: JEOL JSM-IT800 Manual §6.1, p.94]",
            "Inspect EDS detector cooling lines for ice/condensation "
            "[Source: JEOL Service Bulletin SB-2024-07 §2]",
            "Schedule annual PM if usage > 300h since last service "
            "[Source: Mines SIF Maintenance Plan §4]",
        ],
        to_override=to_override,
    )


def scenario_workorder_postrun(to_override: str | None) -> None:
    _hr("Email 3 — Work order: post-run anomaly (detector saturation)  (purple)")
    _send_workorder(
        code="WO-002",
        instrument="JEOL JSM-IT800 SEM-EDS",
        location="Brown Hall B12",
        usage_hours=192.0,
        interval_hours=400.0,
        last_calibrated="2026-03-08",
        issue_type="Detector saturation (DET-SAT)",
        severity="Critical",
        triggered_by="Agent 5 · post-run report BKG-2026-0140",
        anomaly="Detector saturation on consecutive scans at 40 mA tube current.",
        actions=[
            "Reduce tube current to ≤30 mA pending detector service "
            "[Source: Maintenance log DET-SAT-2025-03-12]",
            "Re-zero detector bias and inspect window seal "
            "[Source: JEOL JSM-IT800 Manual §6.2, p.97]",
            "Block bookings until detector service ticket resolved "
            "[Source: Mines SIF Maintenance Plan §3.3]",
        ],
        to_override=to_override,
    )


# ----------------------------------------------------------------------------
# Scenario 4 — Email 4 · Monthly utilization report
# ----------------------------------------------------------------------------
def scenario_monthly(to_override: str | None) -> None:
    _hr("Email 4 — Monthly utilization report  (green)")
    if to_override:
        orig_chair = settings.lab_email_chair
        orig_tech = settings.lab_email_tech
        settings.lab_email_chair = to_override
        settings.lab_email_tech = to_override
    try:
        prev_month = local_now().replace(day=1) - timedelta(days=1)
        period = prev_month.strftime("%B %Y")
        result = email_service.send_monthly_report_email(
            period=period,
            total_bookings=42,
            sops_generated=38,
            avg_fit="86",
            open_work_orders=2,
            utilization=[
                ("JEOL JSM-IT800 SEM-EDS", 92),
                ("Bruker D8 XRD", 64),
                ("Agilent 7900 ICP-MS", 41),
                ("MTS Rock Mechanics Rig", 29),
                ("High-Temp Tube Furnace", 14),
            ],
            insights=[
                "ICP-MS dropped 16% MoM — likely Geochem field-session week (Apr 8–12); pattern matches Q3 2025.",
                "SEM-EDS at 92% for three consecutive months; recommend evaluating an evening shift slot or expanding SEM-Operator certification.",
                "MetEng-Lab accounted for 42% of weekly hours in W17 — equity flag triggered (>40% threshold).",
                "Mean fit score 86 stable vs March (87) — no agent quality regression detected.",
            ],
            action_title="Equity follow-up recommended",
            action_text=(
                "MetEng-Lab exceeded the 40% concentration threshold for one week. "
                "Suggest the chair raise this at the next SIF steering committee meeting."
            ),
        )
        _summarize(result, f"Email 4 dispatched (period: {period})")
    finally:
        if to_override:
            settings.lab_email_chair = orig_chair
            settings.lab_email_tech = orig_tech


# ----------------------------------------------------------------------------
# Dispatcher
# ----------------------------------------------------------------------------
SCENARIOS: dict[str, callable] = {
    "booking": scenario_booking,
    "hitl-training": scenario_hitl_training,
    "hitl-hazmat": scenario_hitl_hazmat,
    "hitl-confidence": scenario_hitl_confidence,
    "workorder-preemptive": scenario_workorder_preemptive,
    "workorder-postrun": scenario_workorder_postrun,
    "monthly": scenario_monthly,
}


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("scenario", nargs="?", help="Name of scenario to fire (or 'all').")
    p.add_argument("--to", dest="to_override", default=os.environ.get("DEMO_EMAIL_TO"),
                   help="Override every recipient. Env: DEMO_EMAIL_TO.")
    p.add_argument("--list", action="store_true", help="List scenarios and exit.")
    args = p.parse_args()

    if args.list or not args.scenario:
        print("Available scenarios (run as: python scripts/demo_emails.py <name>):")
        for k in SCENARIOS:
            print(f"  • {k}")
        print("  • all")
        return 0

    print(f"Resolved local timezone : {LOCAL_TZ_NAME} ({local_tz_label()})")
    print(f"Override recipient      : {args.to_override or '(none — use real recipients)'}")

    if args.scenario == "all":
        for fn in SCENARIOS.values():
            fn(args.to_override)
        return 0

    fn = SCENARIOS.get(args.scenario)
    if not fn:
        print(f"Unknown scenario: {args.scenario}")
        print("Run with --list to see options.")
        return 2
    fn(args.to_override)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
